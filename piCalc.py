import subprocess
import sys
import os

# Function to install a package if not already installed
def install_package(package):
    try:
        __import__(package)
    except ImportError:
        print(f"Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package])

# Function to clear the console
def clear_console():
    if os.name == 'nt':  # For Windows
        os.system('cls')
    else:  # For Linux/macOS
        os.system('clear')

# Ensure required packages are installed
install_package("mpmath")
install_package("python-docx")
install_package("tqdm")

# Clear the console after installing dependencies
clear_console()

from mpmath import mp
from docx import Document
from tqdm import tqdm
import time

mp.dps = 1_000_000

for _ in tqdm(range(1), desc="Calculating Pi to 1 million digits"):
    time.sleep(0.1)  # Simulate progress bar update
    pi_value = str(mp.pi)

doc = Document()
doc.add_heading('Pi to 1 Million Digits', level=1)
doc.add_paragraph(pi_value)

# Save the file in the current working directory
output_file = os.path.join(os.getcwd(), "Pi_to_1_Million_Digits.docx")
doc.save(output_file)

print(f"Pi to 1 million digits has been saved to {output_file}")
